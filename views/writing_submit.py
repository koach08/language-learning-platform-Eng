import streamlit as st
from utils.auth import get_current_user, require_auth
from utils.writing_eval import evaluate_writing, format_writing_feedback
from utils.writing_eval import evaluate_translation, format_translation_feedback
from utils.database import (
    save_writing_submission,
    save_translation_check,
    get_writing_history,
    get_writing_assignments,
    create_writing_assignment,
    get_course_assignments,
    get_assignment_submissions,
    create_submission,
    log_practice,
)

@require_auth
def show():
    user = get_current_user()
    
    st.markdown("## ✍️ ライティング / Writing")
    
    if st.button("← ホームに戻る / Back to Home"):
        if user['role'] == 'teacher':
            st.session_state['current_view'] = 'teacher_home'
        else:
            st.session_state['current_view'] = 'student_home'
        st.rerun()
    
    st.markdown("---")
    
    if user['role'] == 'teacher':
        show_teacher_view()
    else:
        show_student_view()


def show_teacher_view():
    """教員用：課題作成・設定"""
    st.markdown("### 📝 ライティング課題を作成 / Create Writing Assignment")
    
    with st.form("create_writing_assignment"):
        title = st.text_input("課題タイトル / Title", placeholder="例: Unit 1 自己紹介エッセイ")
        
        task_type = st.selectbox("課題タイプ / Type", [
            "自由作文 / Free Writing",
            "エッセイ（意見文） / Essay",
            "メール・手紙 / Email/Letter",
            "要約 / Summary",
            "説明文 / Explanation",
            "スピーチ原稿 / Speech Script",
            "プレゼン原稿 / Presentation Script"
        ])
        
        instructions = st.text_area(
            "指示文 / Instructions",
            placeholder="課題の詳細な指示を入力...",
            height=100
        )
        
        col1, col2 = st.columns(2)
        with col1:
            min_words = st.number_input("最低語数 / Min words", min_value=0, value=100)
        with col2:
            max_words = st.number_input("最大語数 / Max words", min_value=0, value=300)
        
        # 課題として提出必須かどうか
        is_required = st.checkbox("課題として提出必須 / Required submission", value=True)
        st.caption("チェックを外すと、学生の自由練習用になります / Uncheck to make it optional practice")
        
        due_date = st.date_input("提出期限 / Due date")
        
        submitted = st.form_submit_button("課題を作成 / Create", type="primary")
        
        if submitted:
            if title:
                # --- Supabaseに保存 ---
                course_id = st.session_state.get('current_course', {}).get('id')
                if course_id:
                    try:
                        create_writing_assignment(
                            course_id=course_id,
                            title=title,
                            task_type=task_type.split("/")[0].strip(),
                            instructions=instructions,
                            min_words=min_words,
                            max_words=max_words,
                            is_required=is_required,
                            due_date=due_date.isoformat() if due_date else None,
                        )
                        st.success(f"課題「{title}」を作成しました！ / Assignment created!")
                    except Exception as e:
                        st.warning(f"DB保存エラー: {e}")
                        st.success(f"課題「{title}」を作成しました！（ローカルのみ）")
                else:
                    st.success(f"課題「{title}」を作成しました！（※コース未選択のためローカルのみ）")
            else:
                st.error("タイトルを入力してください / Please enter a title")
    
    st.markdown("---")
    st.markdown("### 📋 作成済みの課題 / Created Assignments")
    
    # --- Supabaseから課題一覧を取得 ---
    course_id = st.session_state.get('current_course', {}).get('id')
    assignments = []
    if course_id:
        try:
            assignments = get_writing_assignments(course_id)
        except Exception:
            pass
    
    if assignments:
        for a in assignments:
            config = a.get('config', {}) or {}
            with st.expander(f"📝 {a['title']} (期限: {a.get('due_date', '-')})"):
                st.write(f"**タイプ:** {config.get('task_type', '-')}")
                st.write(f"**語数:** {config.get('min_words', 0)}〜{config.get('max_words', 0)}")
                st.write(f"**指示:** {a.get('instructions', '-')}")
                # 提出数を取得
                try:
                    subs = get_assignment_submissions(a['id'])
                    st.caption(f"提出数: {len(subs)}件")
                except Exception:
                    pass
    else:
        st.info("まだ課題がありません / No assignments yet")


def show_student_view():
    """学生用：課題提出・練習"""
    
    tab1, tab2, tab3, tab4 = st.tabs([
        "📤 課題提出 / Submit Assignment",
        "✏️ 練習添削 / Practice",
        "🔄 翻訳チェック / Translation Check",
        "📊 学習記録 / History"
    ])
    
    with tab1:
        show_assignment_submission()
    
    with tab2:
        show_practice_mode()
    
    with tab3:
        show_translation_check()
    
    with tab4:
        show_writing_history()


def show_assignment_submission():
    """課題提出モード（実データ優先、フォールバックでデモデータ）"""
    
    st.markdown("### 📋 課題一覧 / Assignments")
    
    # --- Supabaseから実課題を取得 ---
    user = get_current_user()
    
    # course_id を複数ソースから取得
    course_id = st.session_state.get('submit_course_id')
    if not course_id:
        course_id = st.session_state.get('current_course', {}).get('id')
    if not course_id:
        try:
            from utils.database import get_student_enrollments
            enrollments = get_student_enrollments(user['id'])
            if enrollments:
                course = enrollments[0].get('courses')
                if course:
                    course_id = course.get('id')
        except Exception:
            pass
    if not course_id:
        course_id = user.get('class_key')
    
    real_assignments = []
    is_demo = False
    
    if course_id:
        try:
            raw = get_writing_assignments(course_id)
            for a in raw:
                config = a.get('config', {}) or {}
                real_assignments.append({
                    "id": a['id'],
                    "title": a.get('title', 'Untitled'),
                    "type": config.get('task_type', 'free_writing'),
                    "instructions": a.get('instructions', ''),
                    "min_words": config.get('min_words', 0),
                    "max_words": config.get('max_words', 500),
                    "due": a.get('due_date', '-') or '-',
                    "required": config.get('is_required', True),
                    "db_id": a['id'],  # Supabaseの実ID
                })
        except Exception:
            pass
    
    # assignmentsテーブルから全課題確認（writing固有フィルタで取れない場合のフォールバック）
    if not real_assignments and course_id:
        try:
            from utils.database import get_student_assignment_status
            all_a = get_student_assignment_status(user['id'], course_id)
            for a in all_a:
                a_type = a.get('type', '')
                if not a_type or a_type.startswith('writing'):
                    real_assignments.append({
                        "id": a.get('assignment_id', ''),
                        "title": a.get('title', 'Untitled'),
                        "type": (a_type.replace('writing_', '') if a_type.startswith('writing_') else 'free_writing'),
                        "instructions": '',
                        "min_words": 0,
                        "max_words": 500,
                        "due": a.get('due_date', '-') or '-',
                        "required": True,
                        "db_id": a.get('assignment_id'),
                    })
        except Exception:
            pass
    
    # assignmentsテーブルから全課題確認（フォールバック）
    if not real_assignments and course_id:
        try:
            from utils.database import get_student_assignment_status
            all_a = get_student_assignment_status(user['id'], course_id)
            for a in all_a:
                a_type = a.get('type', '')
                if not a_type or a_type.startswith('writing'):
                    real_assignments.append({
                        "id": a.get('assignment_id', ''),
                        "title": a.get('title', 'Untitled'),
                        "type": (a_type.replace('writing_', '') if a_type.startswith('writing_') else 'free_writing'),
                        "instructions": '',
                        "min_words": 0,
                        "max_words": 500,
                        "due": a.get('due_date', '-') or '-',
                        "required": True,
                        "db_id": a.get('assignment_id'),
                    })
        except Exception:
            pass

    if real_assignments:
        assignments = real_assignments
    else:
        # デモデータ（サンプルガイド用）
        is_demo = True
        assignments = [
            {
                "id": "demo_1",
                "title": "自己紹介エッセイ / Self-introduction Essay",
                "type": "essay",
                "instructions": "自分自身について紹介するエッセイを書いてください。/ Write an essay introducing yourself.",
                "min_words": 100,
                "max_words": 200,
                "due": "2026/5/15",
                "required": True,
                "db_id": None,
            },
            {
                "id": "demo_2",
                "title": "スピーチ原稿 / Speech Script",
                "type": "speech",
                "instructions": "3分間のスピーチ原稿を書いてください。トピックは自由です。/ Write a 3-minute speech script on any topic.",
                "min_words": 200,
                "max_words": 400,
                "due": "2026/5/22",
                "required": True,
                "db_id": None,
            },
            {
                "id": "demo_3",
                "title": "メール練習（任意）/ Email Practice (Optional)",
                "type": "email",
                "instructions": "教授に質問するメールを書いてください。/ Write an email to ask your professor a question.",
                "min_words": 50,
                "max_words": 150,
                "due": "-",
                "required": False,
                "db_id": None,
            },
        ]
    
    if is_demo:
        st.info("📖 まだ課題が登録されていません。以下はサンプル課題です。\n\nNo assignments registered yet. Below are sample assignments for reference.")
    
    # student_homeから遷移した場合、該当課題を自動選択
    preselect_id = st.session_state.pop('submit_assignment_id', None)
    default_index = 0
    if preselect_id:
        for idx, a in enumerate(assignments):
            if a.get('db_id') == preselect_id or a.get('id') == preselect_id:
                default_index = idx
                break
    
    selected = st.selectbox(
        "課題を選択 / Select assignment",
        assignments,
        index=default_index,
        format_func=lambda x: f"{'📌' if x['required'] else '📝'} {x['title']} (期限: {x['due']})"
    )
    
    if selected:
        st.markdown("---")
        st.markdown(f"### {'📌' if selected['required'] else '📝'} {selected['title']}")
        
        if is_demo:
            st.caption("📖 サンプル課題 / Sample assignment — 提出は記録されますが課題としては未連携です")
        
        if not selected['required']:
            st.caption("🔓 This is optional practice / これは任意の練習です")
        
        st.info(selected['instructions'])
        st.caption(f"語数 / Words: {selected['min_words']}〜{selected['max_words']} | 期限 / Due: {selected['due']}")
        
        st.markdown("---")
        
        # 入力方法選択
        input_method = st.radio(
            "入力方法 / Input Method",
            ["✏️ 直接入力 / Type directly", "📁 ファイルアップロード / Upload file"],
            horizontal=True,
            key="writing_input_method"
        )
        
        text = ""
        uploaded_filename = ""
        
        if input_method == "✏️ 直接入力 / Type directly":
            text = st.text_area(
                "あなたの文章 / Your writing",
                placeholder="Write in English...",
                height=250,
                key="assignment_text"
            )
        
        else:
            st.markdown("**対応形式 / Supported:** Word (.docx), テキスト (.txt), PDF (.pdf)")
            uploaded_file = st.file_uploader(
                "ファイルを選択 / Choose file",
                type=['docx', 'txt', 'pdf'],
                key="writing_file_upload"
            )
            
            if uploaded_file:
                uploaded_filename = uploaded_file.name
                st.caption(f"📄 {uploaded_filename}")
                
                try:
                    if uploaded_file.name.endswith('.txt'):
                        text = uploaded_file.read().decode('utf-8')
                    
                    elif uploaded_file.name.endswith('.docx'):
                        try:
                            from docx import Document
                            import io
                            doc = Document(io.BytesIO(uploaded_file.read()))
                            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                        except ImportError:
                            st.error("Word読込に python-docx が必要です。管理者に連絡してください。")
                            text = ""
                    
                    elif uploaded_file.name.endswith('.pdf'):
                        try:
                            import fitz  # PyMuPDF
                            import io
                            pdf_bytes = uploaded_file.read()
                            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                            text = ""
                            for page in doc:
                                text += page.get_text()
                            doc.close()
                        except ImportError:
                            try:
                                from PyPDF2 import PdfReader
                                import io
                                reader = PdfReader(io.BytesIO(uploaded_file.read()))
                                text = ""
                                for page in reader.pages:
                                    text += page.extract_text() or ""
                            except ImportError:
                                st.error("PDF読込には PyMuPDF または PyPDF2 が必要です。")
                                text = ""
                    
                    if text:
                        st.success(f"✅ ファイル読込完了 / File loaded successfully")
                        with st.expander("📄 読み込んだ内容を確認 / Preview extracted text"):
                            st.text_area("内容", value=text, height=200, disabled=True, key="preview_text")
                
                except Exception as e:
                    st.error(f"ファイル読込エラー / File read error: {e}")
        
        if text:
            word_count = len(text.split())
            
            if selected['min_words'] > 0 and word_count < selected['min_words']:
                st.warning(f"⚠️ 語数 / Words: {word_count} / 最低 Min: {selected['min_words']}")
            elif selected['max_words'] > 0 and word_count > selected['max_words']:
                st.warning(f"⚠️ 語数 / Words: {word_count} / 最大 Max: {selected['max_words']} を超えています")
            else:
                st.success(f"✅ 語数 / Words: {word_count} ({selected['min_words']}〜{selected['max_words']})")
            
            if uploaded_filename:
                st.caption(f"📎 提出ファイル: {uploaded_filename}")
            
            if st.button("📤 提出して評価 / Submit & Evaluate", type="primary"):
                with st.spinner("評価中... / Evaluating... (GPT-4o)"):
                    result = evaluate_writing(
                        text,
                        task_type=selected['type'],
                        level="B1",
                        is_practice=False
                    )
                
                if result.get("success"):
                    show_evaluation_result(result)
                    
                    # --- Supabaseに保存 ---
                    try:
                        save_writing_submission(
                            student_id=user['id'],
                            assignment_id=selected.get('db_id'),  # 実課題ならID、デモならNone
                            text=text,
                            task_type=selected['type'],
                            word_count=len(text.split()),
                            scores=result.get('scores', {}),
                            feedback=format_writing_feedback(result, show_full=True),
                            cefr_level=result.get('cefr_level', ''),
                            is_practice=is_demo,  # デモ課題なら練習扱い
                            course_id=course_id,
                        )
                        if not is_demo:
                            st.success("✅ 提出が記録されました / Submission saved")
                    except Exception as e:
                        st.caption(f"⚠️ DB保存エラー: {e}")
                else:
                    st.error(f"Error: {result.get('error')}")


def show_practice_mode():
    """練習添削モード"""
    
    st.markdown("### ✏️ 練習添削 / Practice & Feedback")
    st.caption("何度でも練習できます / Practice as many times as you like")
    
    practice_type = st.selectbox("タイプ / Type", [
        "自由作文 / Free Writing",
        "エッセイ（意見文） / Essay",
        "メール・手紙 / Email/Letter",
        "スピーチ原稿 / Speech Script",
        "プレゼン原稿 / Presentation Script",
        "要約 / Summary"
    ])
    
    with st.expander("💡 トピック例 / Topic Ideas"):
        topics = {
            "自由作文 / Free Writing": ["My weekend / 週末のこと", "A memorable experience / 思い出", "My hometown / 地元について"],
            "エッセイ（意見文） / Essay": ["Online vs face-to-face classes / オンラインと対面授業", "AI in education / AIと教育", "Social media impact / SNSの影響"],
            "メール・手紙 / Email/Letter": ["Asking a professor / 教授への質問", "Job inquiry / 仕事の問い合わせ", "Thank you letter / お礼の手紙"],
            "スピーチ原稿 / Speech Script": ["Self-introduction / 自己紹介", "My passion / 情熱を持っていること", "A problem I want to solve / 解決したい問題"],
            "プレゼン原稿 / Presentation Script": ["Research summary / 研究の要約", "Product pitch / 商品紹介", "Cultural introduction / 文化紹介"],
            "要約 / Summary": ["News article / ニュース記事", "Book chapter / 本の章", "Lecture notes / 講義ノート"]
        }
        for topic in topics.get(practice_type, []):
            st.markdown(f"- {topic}")
    
    st.markdown("---")
    
    # 入力方法選択
    practice_input = st.radio(
        "入力方法 / Input Method",
        ["✏️ 直接入力 / Type directly", "📁 ファイルアップロード / Upload file"],
        horizontal=True,
        key="practice_input_method"
    )
    
    text = ""
    
    if practice_input == "✏️ 直接入力 / Type directly":
        text = st.text_area(
            "あなたの文章 / Your writing",
            placeholder="Write in English...",
            height=200,
            key="practice_text"
        )
    else:
        st.markdown("**対応形式 / Supported:** Word (.docx), テキスト (.txt), PDF (.pdf)")
        practice_file = st.file_uploader(
            "ファイルを選択 / Choose file",
            type=['docx', 'txt', 'pdf'],
            key="practice_file_upload"
        )
        
        if practice_file:
            try:
                if practice_file.name.endswith('.txt'):
                    text = practice_file.read().decode('utf-8')
                elif practice_file.name.endswith('.docx'):
                    try:
                        from docx import Document
                        import io
                        doc = Document(io.BytesIO(practice_file.read()))
                        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    except ImportError:
                        st.warning("Word読込に python-docx が必要です")
                elif practice_file.name.endswith('.pdf'):
                    try:
                        import fitz
                        import io
                        doc = fitz.open(stream=practice_file.read(), filetype="pdf")
                        text = "".join([page.get_text() for page in doc])
                        doc.close()
                    except ImportError:
                        try:
                            from PyPDF2 import PdfReader
                            import io
                            reader = PdfReader(io.BytesIO(practice_file.read()))
                            text = "".join([page.extract_text() or "" for page in reader.pages])
                        except ImportError:
                            st.warning("PDF読込にはPyMuPDFまたはPyPDF2が必要です")
                
                if text:
                    st.success("✅ ファイル読込完了")
                    with st.expander("📄 内容を確認"):
                        st.text_area("内容", value=text, height=150, disabled=True, key="practice_preview")
            except Exception as e:
                st.error(f"読込エラー: {e}")
    
    if text:
        word_count = len(text.split())
        st.caption(f"📊 語数 / Words: {word_count}")
        
        if word_count < 10:
            st.warning("もう少し書いてください / Please write more (min 10 words)")
        else:
            if st.button("✏️ 添削する / Get Feedback", type="primary"):
                with st.spinner("添削中... / Checking..."):
                    result = evaluate_writing(
                        text,
                        task_type=practice_type.split("/")[0].strip().lower(),
                        level="B1",
                        is_practice=True
                    )
                
                if result.get("success"):
                    show_evaluation_result(result, is_practice=True)
                    
                    # --- Supabaseに保存 ---
                    user = get_current_user()
                    try:
                        save_writing_submission(
                            student_id=user['id'],
                            text=text,
                            task_type=practice_type.split("/")[0].strip(),
                            word_count=len(text.split()),
                            scores=result.get('scores', {}),
                            feedback=format_writing_feedback(result, show_full=False),
                            cefr_level=result.get('cefr_level', ''),
                            is_practice=True,
                            course_id=st.session_state.get('current_course', {}).get('id'),
                        )
                    except Exception as e:
                        st.caption(f"⚠️ DB保存エラー: {e}")
                else:
                    st.error(f"Error: {result.get('error')}")


def show_translation_check():
    """日本語→英語翻訳チェック"""
    
    st.markdown("### 🔄 Translation Check / 翻訳チェック")
    st.markdown("""
    日本語を英語に翻訳する練習です。直訳の問題点と自然な英語を対比して学べます。
    
    Practice translating Japanese to English. Learn by comparing direct translations with natural English.
    """)
    
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 🇯🇵 日本語 / Japanese")
        japanese_text = st.text_area(
            "日本語の原文",
            placeholder="日本語を入力してください...\n例: 私の趣味は映画を見ることです。週末はよく友達と一緒に映画館に行きます。",
            height=200,
            key="japanese_input",
            label_visibility="collapsed"
        )
    
    with col2:
        st.markdown("#### 🇬🇧 Your English Translation / あなたの英訳")
        english_text = st.text_area(
            "英訳",
            placeholder="Translate into English...\nExample: My hobby is watching movies. I often go to the movie theater with my friends on weekends.",
            height=200,
            key="english_input",
            label_visibility="collapsed"
        )
    
    if japanese_text and english_text:
        jp_chars = len(japanese_text)
        en_words = len(english_text.split())
        st.caption(f"📊 日本語: {jp_chars}文字 | English: {en_words} words")
        
        if en_words < 5:
            st.warning("もう少し英語を書いてください / Please write more English")
        else:
            if st.button("🔍 翻訳をチェック / Check Translation", type="primary"):
                with st.spinner("分析中... / Analyzing..."):
                    result = evaluate_translation(
                        japanese_text,
                        english_text,
                        level="B1"
                    )
                
                if result.get("success"):
                    show_translation_result(result)
                    
                    # --- Supabaseに保存 ---
                    user = get_current_user()
                    try:
                        save_translation_check(
                            student_id=user['id'],
                            japanese_text=japanese_text,
                            english_text=english_text,
                            scores=result.get('scores', {}),
                            feedback=format_translation_feedback(result),
                            course_id=st.session_state.get('current_course', {}).get('id'),
                        )
                    except Exception as e:
                        st.caption(f"⚠️ DB保存エラー: {e}")
                else:
                    st.error(f"Error: {result.get('error')}")
    
    # 例文
    with st.expander("💡 練習用の例文 / Practice Examples"):
        examples = [
            ("私は大学で経済学を専攻しています。", "I major in economics at university."),
            ("日本では電車がとても正確です。", "In Japan, trains are very punctual."),
            ("私の夢は世界中を旅行することです。", "My dream is to travel around the world."),
            ("昨日、友達と美味しいラーメンを食べました。", "Yesterday, I ate delicious ramen with my friend."),
            ("将来は国際的な仕事がしたいです。", "I want to work in an international job in the future."),
        ]
        
        for jp, en in examples:
            st.markdown(f"**🇯🇵** {jp}")
            st.markdown(f"**🇬🇧** {en}")
            st.markdown("---")


def show_evaluation_result(result, is_practice=False):
    """評価結果を表示"""
    
    st.success("✅ 評価完了！ / Evaluation Complete!")
    
    scores = result.get("scores", {})
    
    st.markdown("### 📊 Scores / スコア")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Overall / 総合", f"{scores.get('overall', 0)}/100")
    with col2:
        st.metric("CEFR Level", result.get('cefr_level', 'B1'))
    with col3:
        st.metric("Words / 語数", result.get('word_count', 0))
    
    st.markdown("---")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.metric("Grammar / 文法", f"{scores.get('grammar', 0)}")
    with col2:
        st.metric("Vocabulary / 語彙", f"{scores.get('vocabulary', 0)}")
    with col3:
        st.metric("Organization / 構成", f"{scores.get('organization', 0)}")
    with col4:
        st.metric("Content / 内容", f"{scores.get('content', 0)}")
    with col5:
        st.metric("Expression / 表現", f"{scores.get('expression', 0)}")
    
    st.markdown("---")
    feedback = format_writing_feedback(result, show_full=not is_practice)
    st.markdown(feedback)
    
    if is_practice:
        st.markdown("---")
        st.info("💡 フィードバックを参考に、もう一度書いてみましょう！ / Try writing again with this feedback!")


def show_translation_result(result):
    """翻訳評価結果を表示"""
    
    st.success("✅ 分析完了！ / Analysis Complete!")
    
    scores = result.get("scores", {})
    
    st.markdown("### 📊 Scores / スコア")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Overall / 総合", f"{scores.get('overall', 0)}/100")
    with col2:
        st.metric("Accuracy / 正確さ", f"{scores.get('accuracy', 0)}/100")
    with col3:
        st.metric("Naturalness / 自然さ", f"{scores.get('naturalness', 0)}/100")
    with col4:
        st.metric("Grammar / 文法", f"{scores.get('grammar', 0)}/100")
    
    st.markdown("---")
    feedback = format_translation_feedback(result)
    st.markdown(feedback)


def show_writing_history():
    """学習記録を表示"""
    
    st.markdown("### 📊 ライティング学習記録 / Writing History")
    
    user = get_current_user()
    
    # --- Supabaseから取得 ---
    history = []
    try:
        history = get_writing_history(user['id'], limit=30)
    except Exception as e:
        st.warning(f"DB読み込みエラー: {e}")
    
    if not history:
        st.info("まだ学習記録がありません / No history yet")
        return
    
    # 統計
    writing_items = [h for h in history if h.get('content_type') == 'writing']
    translation_items = [h for h in history if h.get('content_type') == 'translation']
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("総提出数", f"{len(history)}件")
    with col2:
        scores = [h.get('score', 0) for h in history if h.get('score')]
        avg = sum(scores) / len(scores) if scores else 0
        st.metric("平均スコア", f"{avg:.1f}点")
    with col3:
        st.metric("ライティング", f"{len(writing_items)}件")
    with col4:
        st.metric("翻訳チェック", f"{len(translation_items)}件")
    
    st.markdown("---")
    
    # 履歴一覧
    for h in history[:15]:
        content_type = h.get('content_type', 'writing')
        details = h.get('details') or {}
        scores_detail = h.get('scores_detail') or {}
        score = h.get('score', 0)
        submitted = h.get('submitted_at', '')[:16].replace('T', ' ')
        is_practice = details.get('is_practice', False)
        
        if content_type == 'translation':
            label = "🔄 翻訳チェック"
        elif is_practice:
            label = f"✏️ 練習 ({details.get('task_type', '-')})"
        else:
            label = f"📤 課題提出 ({details.get('task_type', '-')})"
        
        with st.expander(f"{label} — {score}点 — {submitted}"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("総合", f"{score}点")
            with col2:
                st.metric("語彙", f"{scores_detail.get('vocabulary', '-')}")
            with col3:
                st.metric("文法", f"{scores_detail.get('grammar', '-')}")
            
            if details.get('cefr_level'):
                st.caption(f"CEFR: {details['cefr_level']} | 語数: {details.get('word_count', 0)}")
            
            text_preview = h.get('student_text', '')
            if text_preview:
                st.text_area("提出テキスト", text_preview[:300], height=80, disabled=True, key=f"hist_{h.get('id', submitted)}")
